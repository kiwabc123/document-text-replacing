#!/usr/bin/env python3
from docx import Document

doc = Document('C:\\Users\\tiradet.j\\Downloads\\basic-invoice (1) (2).docx')
print('=== PARAGRAPHS ===')
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'Para {i}: {para.text[:100]}')

print('\n=== TABLES ===')
for table_idx, table in enumerate(doc.tables):
    print(f'Table {table_idx}:')
    for row_idx, row in enumerate(table.rows):
        row_text = ' | '.join([cell.text for cell in row.cells])
        if row_text.strip():
            print(f'  Row {row_idx}: {row_text}')
