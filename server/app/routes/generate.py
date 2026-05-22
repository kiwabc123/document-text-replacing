"""
PDF generation route - handles PDF generation from templates using LibreOffice
"""

from flask import Blueprint, request, jsonify
import os
import logging
import subprocess
import tempfile
import base64
from pathlib import Path

logger = logging.getLogger(__name__)
bp = Blueprint('generate', __name__, url_prefix='/api/pdf')

def fill_docx_variables(docx_path, variables, table_data=None):
    """Fill DOCX template with variables and table data"""
    try:
        from docx import Document
        from docx.shared import Pt
        import re
        
        doc = Document(docx_path)
        
        # Track replacement statistics
        vars_replaced = set()
        para_count = 0
        para_modified = 0
        
        # Replace variables in paragraphs
        for paragraph in doc.paragraphs:
            # Collect all text from runs
            full_text = ''.join([run.text for run in paragraph.runs])
            para_count += 1
            
            # Replace variables
            original_text = full_text
            for var_name, var_value in variables.items():
                pattern = f"{{{{{var_name}}}}}"
                if pattern in full_text:
                    vars_replaced.add(var_name)
                    full_text = full_text.replace(pattern, str(var_value))
            
            # Only modify if text changed
            if full_text != original_text:
                para_modified += 1
                logger.info(f"Para modified: {original_text[:100]} -> {full_text[:100]}")
                # Clear paragraph content properly - use list() to avoid live list issues
                for run in list(paragraph.runs):
                    # Remove the run element from the paragraph
                    run._element.getparent().remove(run._element)
                # Add new run with replaced text
                paragraph.add_run(full_text)
        
        logger.info(f"Paragraphs: {para_count} total, {para_modified} modified. Variables replaced in paragraphs: {vars_replaced}")
        
        # Replace variables in tables and fill table data
        if table_data:
            table_vars_replaced = set()
            for table_idx, table in enumerate(doc.tables):
                if len(table.rows) < 1:
                    continue
                
                # First, replace variables in existing rows
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text = ''.join([run.text for run in paragraph.runs])
                            original_text = full_text
                            for var_name, var_value in variables.items():
                                pattern = f"{{{{{var_name}}}}}"
                                if pattern in full_text:
                                    table_vars_replaced.add(var_name)
                                full_text = full_text.replace(pattern, str(var_value))
                            
                            if full_text != original_text:
                                for run in list(paragraph.runs):
                                    run._element.getparent().remove(run._element)
                                paragraph.add_run(full_text)
            
            logger.info(f"Variables replaced in tables: {table_vars_replaced}")
            
            # Now fill table rows with table data if available
            # Try to find table by name (first table is usually "items" or "table_0")
            table_name = None
            if len(table.rows) > 0:
                # Get column headers to match table
                headers = []
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip().lower())
                
                # Find matching table_data by checking if any key has matching columns
                for data_key, rows in table_data.items():
                    # If this table_data has matching structure, use it
                    if rows and len(rows) > 0:
                        table_name = data_key
                        break
            
            # Fill table rows
            if table_name and table_name in table_data:
                rows_data = table_data[table_name]
                
                # Keep header row (row 0), replace/add data rows
                # Remove existing data rows (keep header)
                while len(table.rows) > 1:
                    tr = table.rows[1]._tr
                    tr.getparent().remove(tr)
                
                # Add new rows with data
                for row_data in rows_data:
                    new_row = table.add_row()
                    
                    # Get column headers
                    headers = []
                    for cell in table.rows[0].cells:
                        headers.append(cell.text.strip().lower().replace(' ', '_'))
                    
                    # Fill cells in the new row
                    for col_idx, header in enumerate(headers):
                        if col_idx < len(new_row.cells):
                            # Find matching key in row_data
                            cell_value = ''
                            for key, value in row_data.items():
                                if key.lower().replace(' ', '_') == header or key == header:
                                    cell_value = str(value)
                                    break
                            
                            # Set cell text
                            new_row.cells[col_idx].text = cell_value
        else:
            # Just replace variables in tables without filling data
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text = ''.join([run.text for run in paragraph.runs])
                            original_text = full_text
                            for var_name, var_value in variables.items():
                                pattern = f"{{{{{var_name}}}}}"
                                full_text = full_text.replace(pattern, str(var_value))
                            
                            if full_text != original_text:
                                for run in list(paragraph.runs):
                                    run._element.getparent().remove(run._element)
                                paragraph.add_run(full_text)
        
        return doc
        
    except Exception as e:
        logger.error(f"Error filling DOCX: {e}")
        raise

def convert_docx_to_pdf_with_libreoffice(docx_path, output_path):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        # Use LibreOffice to convert DOCX to PDF
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_path),
            docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, timeout=30)
        
        if result.returncode != 0:
            logger.error(f"LibreOffice conversion error: {result.stderr.decode()}")
            raise Exception("LibreOffice conversion failed")
        
        # LibreOffice outputs with same name but .pdf extension
        libreoffice_output = docx_path.replace('.docx', '.pdf')
        if os.path.exists(libreoffice_output):
            os.rename(libreoffice_output, output_path)
            return True
        
        return False
        
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice conversion timeout")
        raise
    except Exception as e:
        logger.error(f"Conversion error: {e}")
        raise

@bp.route('/generate', methods=['POST'])
def generate_pdf():
    """Generate PDF from DOCX template with variables filled"""
    try:
        body = request.get_json()
        
        if not body:
            return jsonify({'error': 'No data provided'}), 400
        
        variables = body.get('variables', {})
        table_data = body.get('tableData', {})
        template_content = body.get('templateContent', '')  # base64 encoded
        
        if not template_content:
            return jsonify({'error': 'No template provided'}), 400
        
        logger.info(f"Generating PDF with {len(variables)} variables and {len(table_data)} tables")
        
        # Create temp directory
        temp_dir = tempfile.gettempdir()
        temp_docx = os.path.join(temp_dir, 'temp_template.docx')
        temp_pdf = os.path.join(temp_dir, 'temp_output.pdf')
        
        # Decode base64 and write DOCX
        docx_bytes = base64.b64decode(template_content)
        with open(temp_docx, 'wb') as f:
            f.write(docx_bytes)
        
        # Fill variables and table data
        filled_doc = fill_docx_variables(temp_docx, variables, table_data if table_data else None)
        filled_docx_path = os.path.join(temp_dir, 'temp_filled.docx')
        filled_doc.save(filled_docx_path)
        
        # Save a copy for debugging
        debug_docx_path = os.path.join('/app/uploads', 'filled_output.docx')
        filled_doc.save(debug_docx_path)
        logger.info(f"Filled DOCX saved for debugging: {debug_docx_path}")
        
        # Convert to PDF using LibreOffice
        convert_docx_to_pdf_with_libreoffice(filled_docx_path, temp_pdf)
        
        # Read PDF and encode to base64
        with open(temp_pdf, 'rb') as f:
            pdf_bytes = f.read()
        
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        
        # Cleanup
        for f in [temp_docx, filled_docx_path, temp_pdf]:
            if os.path.exists(f):
                os.remove(f)
        
        logger.info(f"PDF generated successfully: {len(pdf_bytes)} bytes")
        
        return jsonify({
            'success': True,
            'data': {
                'pdf': pdf_base64,
                'size': len(pdf_bytes),
            }
        }), 200
        
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        return jsonify({'error': f'PDF generation failed: {str(e)}'}), 500
