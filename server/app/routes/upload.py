"""
Upload route - handles DOCX file uploads and variable extraction
"""

from flask import Blueprint, request, jsonify, current_app
from werkzeug.utils import secure_filename
import os
import logging
from docx import Document
import tempfile
import re

logger = logging.getLogger(__name__)
bp = Blueprint('upload', __name__, url_prefix='/api/templates')

ALLOWED_EXTENSIONS = {'docx', 'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def infer_variable_type(var_name):
    """Infer variable type based on name"""
    var_lower = var_name.lower()
    if 'email' in var_lower or 'mail' in var_lower:
        return 'email'
    elif 'phone' in var_lower or 'tel' in var_lower:
        return 'phone'
    elif 'date' in var_lower:
        return 'date'
    elif 'amount' in var_lower or 'price' in var_lower or 'cost' in var_lower or 'total' in var_lower:
        return 'currency'
    elif 'qty' in var_lower or 'quantity' in var_lower or 'count' in var_lower:
        return 'number'
    else:
        return 'text'

def extract_tables_from_docx(doc):
    """Extract table schemas from DOCX"""
    tables = []
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 1:
            continue
        
        # Get header row
        header_cells = table.rows[0].cells
        headers = [cell.text.strip() for cell in header_cells]
        
        # Skip if no headers
        if not headers or all(h == '' for h in headers):
            continue
        
        # Create columns from headers
        columns = []
        for header in headers:
            col_id = header.lower().replace(' ', '_')
            columns.append({
                'id': col_id,
                'name': col_id,
                'type': infer_variable_type(col_id),
                'label': header,
                'required': True
            })
        
        # Create table schema
        table_name = f"table_{table_idx}" if table_idx > 0 else "items"
        tables.append({
            'id': table_name,
            'name': table_name.capitalize(),
            'columns': columns
        })
    
    return tables

def convert_to_template_variables(var_names):
    """Convert variable names to TemplateVariable objects"""
    template_vars = []
    for name in var_names:
        # Generate label from variable name
        # Handle spaces, camelCase, and snake_case
        label = name.strip()
        label = re.sub(r'([a-z])([A-Z])', r'\1 \2', label)  # camelCase -> space
        label = re.sub(r'_', ' ', label)  # snake_case -> space
        label = ' '.join(word.capitalize() for word in label.split())  # Title case each word
        
        template_vars.append({
            'name': name,
            'label': label,
            'type': infer_variable_type(name),
            'required': True,
            'placeholder': f'Enter {label.lower()}'
        })
    return template_vars

def extract_docx_variables(docx_path):
    """Extract variables and content from DOCX file"""
    try:
        doc = Document(docx_path)
        variables = set()
        
        # Extract variables from text (format: {{variableName}} - supports spaces)
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find all {{...}} patterns (allows spaces and special chars inside)
            matches = re.findall(r'\{\{([^}]+?)\}\}', text)
            for match in matches:
                # Clean up: strip whitespace and skip empty
                clean_name = match.strip()
                if clean_name:
                    variables.add(clean_name)
        
        # Also check in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    matches = re.findall(r'\{\{([^}]+?)\}\}', text)
                    for match in matches:
                        clean_name = match.strip()
                        if clean_name:
                            variables.add(clean_name)
        
        return {
            'variables': sorted(list(variables)),
            'pages': len(doc.paragraphs),
        }
    except Exception as e:
        logger.error(f"Error extracting DOCX variables: {e}")
        raise

@bp.route('/upload', methods=['POST'])
def upload_template():
    """Upload and parse template file (DOCX or PDF)"""
    try:
        logger.info(f"Upload request received. Files: {list(request.files.keys())}, Content-Type: {request.content_type}")
        
        # Check if file is present
        if 'file' not in request.files:
            logger.error(f"No 'file' field in request. Available fields: {list(request.files.keys())}")
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        logger.info(f"File received: {file.filename}, size: {len(file.getbuffer())} bytes")
        
        if file.filename == '':
            logger.error("File filename is empty")
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            logger.error(f"Invalid file extension: {file.filename}")
            return jsonify({'error': 'Only .docx and .pdf files are supported'}), 400
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        
        logger.info(f"Processing file: {filename}")
        
        # Extract variables based on file type
        if filename.endswith('.docx'):
            result = extract_docx_variables(file_path)
            tables = extract_tables_from_docx(Document(file_path))
            file_type = 'docx'
        else:
            # For PDF, return basic structure for now
            result = {'variables': [], 'pages': 0}
            tables = []
            file_type = 'pdf'
        
        # Convert variable names to TemplateVariable objects
        template_variables = convert_to_template_variables(result['variables'])
        
        # Create response with file data
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        import base64
        import datetime
        import time
        
        response_data = {
            'success': True,
            'data': {
                'id': f"template_{int(time.time() * 1000)}",
                'name': filename.rsplit('.', 1)[0],
                'fileName': filename,
                'uploadedAt': datetime.datetime.now().isoformat(),
                'variables': template_variables,
                'tables': tables,
                'content': base64.b64encode(file_content).decode('utf-8'),
                'templateType': 'docx',
            }
        }
        
        # Cleanup
        os.remove(file_path)
        
        logger.info(f"Template uploaded successfully. Variables: {[v['name'] for v in template_variables]}")
        return jsonify(response_data), 201
        
    except Exception as e:
        logger.error(f"Upload error: {e}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500
