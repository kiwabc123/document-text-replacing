from docx import Document
import re

docx_path = r'C:\Users\tiradet.j\Downloads\basic-invoice (1) (2).docx'

doc = Document(docx_path)

var_pattern = r'\{\{([^}]+?)\}\}'

print("=== SECTIONS ===")
print(f"Total sections: {len(doc.sections)}")
for i, section in enumerate(doc.sections):
    print(f"\nSection {i}:")
    print(f"  Header paragraphs: {len(section.header.paragraphs)}")
    print(f"  Footer paragraphs: {len(section.footer.paragraphs)}")

print("\n=== PARAGRAPHS ===")
print(f"Total paragraphs: {len(doc.paragraphs)}")
para_with_vars = 0
for i, para in enumerate(doc.paragraphs):
    matches = re.findall(var_pattern, para.text)
    if matches:
        para_with_vars += 1
        print(f"Para {i}: {para.text}")

print(f"\n\nTotal paragraphs with variables: {para_with_vars}")

print("\n=== TABLES ===")
print(f"Total tables: {len(doc.tables)}")
for table_idx, table in enumerate(doc.tables):
    print(f"\nTable {table_idx}:")
    print(f"  Rows: {len(table.rows)}, Cols: {len(table.columns)}")
    cells_with_vars = 0
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            matches = re.findall(var_pattern, cell.text)
            if matches:
                cells_with_vars += 1
                print(f"    Cell [{row_idx},{col_idx}]: {cell.text}")
    print(f"  Total cells with variables: {cells_with_vars}")
